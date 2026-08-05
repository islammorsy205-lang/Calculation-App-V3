# bridge_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import ast
import re
import matplotlib
matplotlib.use('Agg') # 💡 وضع الخوادم لمنع أي انهيار للواجهة
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================================================
# 1. Native HTML Parser Engine (Extracting EVERYTHING Safely)
# =========================================================
def parse_bridge_html_native(html_content):
    """يقرأ مصفوفات الجافاسكريبت وجداول الـ HTML بنظام الاستخراج النصي الآمن تماماً لتجنب أعطال الذاكرة"""
    
    # 1. Extract Arrays using pure string finding
    def extract_array_str(var_name):
        start_marker = f"const {var_name}"
        idx_start = html_content.find(start_marker)
        if idx_start == -1: return "[]"
        idx_bracket = html_content.find("[", idx_start)
        idx_end = html_content.find("];", idx_bracket)
        if idx_bracket == -1 or idx_end == -1: return "[]"
        return html_content[idx_bracket : idx_end + 1]

    nodes_raw = extract_array_str("globalNodes")
    elems_raw = extract_array_str("globalElements")
    
    nodes, elements = None, None
    def clean_js_to_python(js_str):
        js_str = re.sub(r'([{,]\s*)([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', js_str)
        js_str = js_str.replace(': false', ': False').replace(': true', ': True')
        js_str = js_str.replace(':false', ':False').replace(':true', ':True')
        return js_str
        
    try:
        nodes = ast.literal_eval(clean_js_to_python(nodes_raw))
        elements = ast.literal_eval(clean_js_to_python(elems_raw))
    except Exception as e:
        st.error(f"⚠️ خطأ في قراءة مصفوفات النموذج: {e}")

    # 2. Extract Tables Natively (Bypassing Pandas to prevent Crashes)
    tables = []
    table_blocks = re.findall(r'<table.*?>(.*?)</table>', html_content, re.DOTALL | re.IGNORECASE)
    
    for block in table_blocks:
        rows = re.findall(r'<tr.*?>(.*?)</tr>', block, re.DOTALL | re.IGNORECASE)
        parsed_table = []
        for row in rows:
            cells = re.findall(r'<(t[hd]).*?>(.*?)</\1>', row, re.DOTALL | re.IGNORECASE)
            clean_cells = [re.sub(r'<.*?>', '', c[1]).strip() for c in cells]
            if clean_cells:
                parsed_table.append(clean_cells)
        if parsed_table:
            tables.append(parsed_table)

    return nodes, elements, tables

# =========================================================
# 2. SAP2000 Plotting Engine (All Diagrams & Styles)
# =========================================================
def safe_render_fig(fig):
    try:
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=200, bbox_inches='tight', transparent=True)
        return buf.getvalue()
    finally:
        plt.close(fig)

def draw_base_structure(ax, nodes, elements):
    nodes_dict = {n['id']: n for n in nodes}
    for el in elements:
        n1, n2 = nodes_dict.get(el['i']), nodes_dict.get(el['j'])
        if not n1 or not n2: continue
        ax.plot([float(n1['x']), float(n2['x'])], [float(n1['y']), float(n2['y'])], color='dimgray', linewidth=1.5, zorder=1)
        
    for n in nodes:
        if n.get('fixX') or n.get('fixY') or n.get('fixT'):
            x, y = float(n['x']), float(n['y'])
            t = 'Fixed' if (n.get('fixX') and n.get('fixY') and n.get('fixT')) else \
                'Hinged' if (n.get('fixX') and n.get('fixY')) else 'Roller'
                
            if t == 'Hinged' or t == 'Fixed':
                h, w = 0.5, 0.4
                p1, p2, p3 = (x, y), (x + w/2, y - h), (x - w/2, y - h)
                ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5))
                ax.plot([x - w, x + w], [y - h, y - h], color='limegreen', lw=2.0, zorder=4)
            elif t == 'Roller':
                h, w, r = 0.4, 0.3, 0.12
                p1, p2, p3 = (x, y), (x + w/2, y - h), (x - w/2, y - h)
                ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5))
                ax.add_patch(plt.Circle((x, y - h - r), r, facecolor='none', edgecolor='limegreen', lw=1.5, zorder=5))
                ax.plot([x - 0.2, x + 0.2], [y - h - 2*r, y - h - 2*r], color='limegreen', lw=2.0, zorder=4)
    return nodes_dict

def draw_sap2000_forces(val_key, nodes, elements, scale, is_axial=False):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    
    nodes_dict = draw_base_structure(ax, nodes, elements)

    for el in elements:
        n1, n2 = nodes_dict.get(el['i']), nodes_dict.get(el['j'])
        if not n1 or not n2: continue
        
        x1, y1 = float(n1['x']), float(n1['y'])
        x2, y2 = float(n2['x']), float(n2['y'])
        dx, dy = x2 - x1, y2 - y1
        L_s = np.hypot(dx, dy)
        if L_s < 1e-5: continue
        
        c, s = dx/L_s, dy/L_s
        diag_arr = el.get('axialDiag' if is_axial else 'diag', [])
        if not diag_arr: continue
        
        ts = np.array([float(pt.get('t', 0)) for pt in diag_arr])
        vals_orig = np.array([float(pt.get('n' if is_axial else val_key.lower(), 0)) for pt in diag_arr])
        if len(vals_orig) == 0: continue
        
        plot_vals = -vals_orig if val_key != 'N' else vals_orig
        
        px_arr = x1 + c * (ts * L_s) - s * plot_vals * scale
        py_arr = y1 + s * (ts * L_s) + c * plot_vals * scale
        
        color_pos, color_neg = 'blue', 'red'
        
        ax.plot([x1, px_arr[0]], [y1, py_arr[0]], color=color_pos if vals_orig[0] >= 0 else color_neg, linewidth=0.8)
        for k in range(len(px_arr)-1):
            avg_v = (vals_orig[k] + vals_orig[k+1]) / 2.0
            seg_color = color_pos if avg_v >= 0 else color_neg
            ax.plot([px_arr[k], px_arr[k+1]], [py_arr[k], py_arr[k+1]], color=seg_color, linewidth=0.8)
        ax.plot([px_arr[-1], x2], [py_arr[-1], y2], color=color_pos if vals_orig[-1] >= 0 else color_neg, linewidth=0.8)
        
        num_lines = max(2, int(L_s / 0.4))
        for i in range(1, num_lines):
            frac = i / num_lines
            lx, ly = x1 + frac * dx, y1 + frac * dy
            idx_val = int(frac * (len(plot_vals)-1))
            lv = plot_vals[idx_val]
            ax.plot([lx, lx - s * lv * scale], [ly, ly + c * lv * scale], color=color_pos if vals_orig[idx_val] >= 0 else color_neg, linewidth=0.3, alpha=0.6)
            
        max_idx = np.argmax(np.abs(vals_orig))
        if abs(vals_orig[max_idx]) > 0.1:
            ax.text(px_arr[max_idx] - s*0.4, py_arr[max_idx] + c*0.4, f"{abs(vals_orig[max_idx]):.1f}", color='black', fontsize=6, ha='center', va='center')

    return safe_render_fig(fig)

def draw_sap2000_deflection(nodes, elements, defl_scale):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    
    nodes_dict = draw_base_structure(ax, nodes, elements)
    max_defl = 0
    max_pt = None
    
    for el in elements:
        n1, n2 = nodes_dict.get(el['i']), nodes_dict.get(el['j'])
        if not n1 or not n2: continue
        x1 = float(n1['x']) + float(n1.get('dx', 0)) * defl_scale
        y1 = float(n1['y']) + float(n1.get('dy', 0)) * defl_scale
        x2 = float(n2['x']) + float(n2.get('dx', 0)) * defl_scale
        y2 = float(n2['y']) + float(n2.get('dy', 0)) * defl_scale
        ax.plot([x1, x2], [y1, y2], color='red', linestyle='--', linewidth=1.5, alpha=0.8, zorder=3)
        
    for n in nodes:
        dx, dy = float(n.get('dx', 0)), float(n.get('dy', 0))
        defl = np.hypot(dx, dy)
        if defl > max_defl:
            max_defl = defl
            max_pt = (float(n['x']) + dx * defl_scale, float(n['y']) + dy * defl_scale)
            
    if max_pt and max_defl > 0.0001:
        ax.annotate(f"Max Defl: {max_defl*1000:.2f} mm", xy=max_pt, xytext=(max_pt[0]+1, max_pt[1]+1),
                    arrowprops=dict(facecolor='red', shrink=0.05, width=1.5, headwidth=6),
                    fontsize=8, color='red', fontweight='bold', zorder=10)

    return safe_render_fig(fig)

def draw_sap2000_reactions(nodes, elements):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    draw_base_structure(ax, nodes, elements)
    
    for n in nodes:
        rx, ry = float(n.get('rx', 0)), float(n.get('ry', 0))
        x, y = float(n['x']), float(n['y'])
        arr_len = 1.2
        if abs(ry) > 0.1:
            dy = -arr_len if ry > 0 else arr_len
            ax.arrow(x, y + dy, 0, -dy*0.8, head_width=0.3, head_length=0.4, fc='darkorange', ec='darkorange', lw=1.5, zorder=6)
            ax.text(x, y + dy - np.sign(dy)*0.3, f"{abs(ry):.1f} kN", color='black', fontsize=7, fontweight='bold', ha='center')
        if abs(rx) > 0.1:
            dx = -arr_len if rx > 0 else arr_len
            ax.arrow(x + dx, y, -dx*0.8, 0, head_width=0.3, head_length=0.4, fc='darkorange', ec='darkorange', lw=1.5, zorder=6)
            ax.text(x + dx - np.sign(dx)*0.3, y, f"{abs(rx):.1f} kN", color='black', fontsize=7, fontweight='bold', va='center')

    return safe_render_fig(fig)

def draw_sap2000_loads(nodes, elements):
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    
    nodes_dict = draw_base_structure(ax, nodes, elements)
    max_w = max([abs(float(el.get('wTotal', 0))) for el in elements] + [1])
    scale_h = 1.5 / max_w
    
    for el in elements:
        w = float(el.get('wTotal', 0))
        if abs(w) < 0.1: continue
        n1, n2 = nodes_dict.get(el['i']), nodes_dict.get(el['j'])
        if not n1 or not n2: continue
        x1, y1 = float(n1['x']), float(n1['y'])
        x2, y2 = float(n2['x']), float(n2['y'])
        h = abs(w) * scale_h
        
        ax.add_patch(Polygon([(x1,y1), (x1, y1+h), (x2, y2+h), (x2, y2)], facecolor='royalblue', edgecolor='blue', alpha=0.3, zorder=2))
        num_arr = max(1, int(np.hypot(x2-x1, y2-y1) / 0.5))
        for i in range(1, num_arr):
            fx, fy = x1 + (x2-x1) * (i/num_arr), y1 + (y2-y1) * (i/num_arr)
            ax.arrow(fx, fy+h, 0, -h*0.8, head_width=0.1, head_length=0.2, fc='blue', ec='blue', lw=0.5, zorder=3)
        ax.text((x1+x2)/2, (y1+y2)/2 + h + 0.3, f"{abs(w):.2f} kN/m", color='blue', fontsize=7, fontweight='bold', ha='center')

    return safe_render_fig(fig)

# =========================================================
# 3. Comprehensive Report Generator (Word)
# =========================================================
def generate_comprehensive_bridge_report(nodes, elements, tables, img_bytes_dict):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_line(text, bold=False, color=None, size=11):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.rtl = False
        if color: r.font.color.rgb = color
        return p

    def add_native_table_to_word(doc, table_data, title):
        if not table_data or len(table_data) < 2: return
        add_line(title, bold=True, size=13, color=RGBColor(0,0,128))
        
        cols_count = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols_count)
        table.style = 'Table Grid'
        
        for r_idx, row_data in enumerate(table_data):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < cols_count:
                    row_cells[c_idx].text = str(cell_text)
                    for paragraph in row_cells[c_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            if r_idx == 0:
                                run.font.bold = True
                            else:
                                text_up = str(cell_text).upper()
                                if "PASS" in text_up or "SAFE" in text_up:
                                    run.font.color.rgb = RGBColor(0, 128, 0)
                                    run.font.bold = True
                                elif "FAIL" in text_up or "UNSAFE" in text_up:
                                    run.font.color.rgb = RGBColor(255, 0, 0)
                                    run.font.bold = True
        doc.add_paragraph()

    # --- 1. Cover ---
    p_title = doc.add_paragraph()
    force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR BRIDGE STRUCTURE (FULL FEA)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    
    add_line("="*60, bold=True)
    add_line("1. Executive Summary:", bold=True, size=14)
    add_line(f"- Total Nodes Analyzed: {len(nodes)}")
    add_line(f"- Total Elements Analyzed: {len(elements)}")
    
    max_m, max_v, max_n_tens, max_n_comp = 0, 0, 0, 0
    for el in elements:
        if 'maxM' in el: max_m = max(max_m, abs(float(el['maxM'])))
        if 'maxV' in el: max_v = max(max_v, abs(float(el['maxV'])))
        if 'maxN' in el:
            n_val = float(el['maxN'])
            if n_val > 0: max_n_tens = max(max_n_tens, n_val)
            else: max_n_comp = min(max_n_comp, n_val)
            
    add_line("\n2. Global Force Extremes:", bold=True)
    add_line(f"- Maximum Bending Moment = {max_m:.2f} kN.m")
    add_line(f"- Maximum Shear Force = {max_v:.2f} kN")
    add_line(f"- Maximum Axial Tension = {max_n_tens:.2f} kN")
    add_line(f"- Maximum Axial Compression = {abs(max_n_comp):.2f} kN")
    
    add_line("\n3. Support Reactions:", bold=True)
    for n in nodes:
        if n.get('fixX') or n.get('fixY') or n.get('fixT'):
            rx, ry = float(n.get('rx', 0)), float(n.get('ry', 0))
            add_line(f"- Node {n.get('name', 'N')}: Rx = {rx:.2f} kN | Ry = {ry:.2f} kN", color=RGBColor(0, 100, 0))

    # --- 2. Tables ---
    doc.add_page_break()
    add_line("4. Structural Analysis Tables & Safety Checks:", bold=True, size=15)
    doc.add_paragraph()
    
    table_titles = [
        "Nodal Displacements", "Element Internal Forces Summary", "BMD Extreme Values",
        "SFD Extreme Values", "Axial Force (Main Members)", "Axial Force (Bracing)",
        "Deflection Check", "Support Reactions Summary", "Applied Loads"
    ]
    
    for i, tbl_data in enumerate(tables):
        title = table_titles[i] if i < len(table_titles) else f"Data Table {i+1}"
        add_native_table_to_word(doc, tbl_data, f"Table {i+1}: {title}")

    # --- 3. Diagrams ---
    doc.add_page_break()
    add_line("5. Structural Diagrams (SAP2000 Convention):", bold=True, size=15)
    
    diagram_order = [
        ('L', "Global Applied Loads (kN/m)"),
        ('N', "Axial Force Diagram (kN) [Blue = Tension, Red = Compression]"),
        ('V', "Shear Force Diagram (kN)"),
        ('M', "Bending Moment Diagram (kN.m)"),
        ('D', "Global Deflection Deformed Shape (mm)"),
        ('R', "Global Support Reactions (kN)")
    ]
    
    for key, title in diagram_order:
        img_bytes = img_bytes_dict.get(key)
        if not img_bytes: continue
        
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(16.5))
        
        p_txt = doc.add_paragraph()
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_txt = p_txt.add_run(title)
        r_txt.font.name = 'Arial'
        r_txt.font.size = Pt(10)
        r_txt.font.bold = True
        r_txt.underline = True
        doc.add_page_break()
    
    out = io.BytesIO()
    doc.save(out)
    return out

# =========================================================
# 4. Main UI Module
# =========================================================
def render_bridge_module():
    st.markdown("## 🌉 Bridge Formwork & Structures (Advanced FEA Extractor)")
    st.info("💡 **Smart Extractor:** Upload your Acrow Bridge HTML Report to extract all calculations, safety checks, and diagrams. Convert them instantly to SAP2000 style and generate a complete Word calculation sheet.")
    
    uploaded_file = st.file_uploader("📂 Upload Acrow Bridge FEA HTML File", type=["html"])
    
    if uploaded_file is not None:
        html_content = uploaded_file.getvalue().decode("utf-8")
        
        with st.spinner("Extracting Nodes, Elements, and Checking Tables safely..."):
            nodes, elements, tables = parse_bridge_html_native(html_content)
            
        if nodes and elements:
            st.success(f"✅ Data Extracted Successfully! Found **{len(nodes)} Nodes**, **{len(elements)} Elements**, and **{len(tables)} Analysis Tables**.")
            st.markdown("---")
            
            diag_placeholder = st.empty()
            
            # Sliders تحت الدايجرامات
            st.markdown("### 🎛️ Customize Diagram Scales")
            c_s1, c_s2, c_s3, c_s4 = st.columns(4)
            sc_n = c_s1.slider("Axial Scale", 0.001, 0.050, 0.015, step=0.001)
            sc_v = c_s2.slider("Shear Scale", 0.001, 0.050, 0.015, step=0.001)
            sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.015, step=0.001)
            sc_d = c_s4.slider("Deflection Scale", 1.0, 100.0, 20.0, step=1.0)
            
            img_bufs = {}
            with diag_placeholder.container():
                st.markdown("### 🎛️ Live SAP2000-Style Diagrams")
                try:
                    img_bufs = {
                        'L': draw_sap2000_loads(nodes, elements),
                        'N': draw_sap2000_forces('N', nodes, elements, sc_n, is_axial=True),
                        'V': draw_sap2000_forces('V', nodes, elements, sc_v),
                        'M': draw_sap2000_forces('M', nodes, elements, sc_m),
                        'D': draw_sap2000_deflection(nodes, elements, sc_d),
                        'R': draw_sap2000_reactions(nodes, elements)
                    }
                    
                    st.image(img_bufs['L'], caption="Applied Loads Diagram")
                    c_p1, c_p2 = st.columns(2)
                    c_p1.image(img_bufs['M'], caption="Bending Moment Diagram (kN.m)")
                    c_p2.image(img_bufs['V'], caption="Shear Force Diagram (kN)")
                    c_p3, c_p4 = st.columns(2)
                    c_p3.image(img_bufs['N'], caption="Axial Force Diagram (kN)")
                    c_p4.image(img_bufs['D'], caption="Deflection Deformed Shape")
                    st.image(img_bufs['R'], caption="Support Reactions Diagram")
                except Exception as e:
                    st.error(f"❌ حدث خطأ أثناء الرسم: {e}")
                
            st.markdown("---")
            
            if img_bufs:
                if st.button("🚀 Process & Generate Calculation Sheet", type="primary", use_container_width=True):
                    with st.spinner("Building Comprehensive Word Document..."):
                        try:
                            docx_out = generate_comprehensive_bridge_report(nodes, elements, tables, img_bufs)
                            st.session_state['bridge_docx_bytes'] = docx_out.getvalue()
                            st.success("✅ Document Ready! You can download it below.")
                        except Exception as e:
                            st.error(f"⚠️ خطأ أثناء تجميع ملف الوورد: {e}")
                
                if 'bridge_docx_bytes' in st.session_state:
                    st.download_button(
                        "⬇️ Download Full Bridge Calculation Sheet (Word)", 
                        data=st.session_state['bridge_docx_bytes'], 
                        file_name="Acrow_Bridge_Full_Report.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
