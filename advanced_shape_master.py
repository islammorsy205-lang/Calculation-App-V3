# advanced_shape_master.py

import streamlit as st
import numpy as np
import pandas as pd
import io
import os
import re
import math
import tempfile
import matplotlib as mpl
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from PIL import Image, ImageChops
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit.components.v1 as components

try:
    import ezdxf
except ImportError:
    st.error("⚠️ مكتبة 'ezdxf' غير موجودة! برجاء كتابة الأمر 'pip install ezdxf' في التيرمينال لتفعيل ميزة قراءة الكاد.")
    ezdxf = None

try:
    from config import SECTIONS_DB, STRUTS_DB
    from report_builder import insert_blue_banner, add_eq, append_pdf_stream_to_word
except ImportError:
    st.error("⚠️ برجاء التأكد من وجود ملفات config.py و report_builder.py")

# =========================================================
# 0. Helper Functions & Styles
# =========================================================
def apply_plot_styles():
    mpl.rcParams['font.family'] = 'sans-serif'
    mpl.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    mpl.rcParams['axes.linewidth'] = 0.3
    mpl.rcParams['font.size'] = 7
    mpl.rcParams['font.weight'] = 'normal'

def get_short_name(sec_name):
    return re.sub(r'\s*\(.*?\)', '', sec_name).strip()

def crop_image_bbox(img_bytes):
    img = Image.open(io.BytesIO(img_bytes)).convert("RGBA")
    bg = Image.new("RGBA", img.size, (255, 255, 255, 0))
    diff = ImageChops.difference(img, bg)
    bbox = diff.getbbox()
    if bbox:
        img = img.crop(bbox)
    out = io.BytesIO()
    img.save(out, format='PNG')
    return out.getvalue()

def safe_render_fig(fig):
    try:
        plt.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=400, bbox_inches='tight', pad_inches=0.0, transparent=True)
        return crop_image_bbox(buf.getvalue())
    finally:
        plt.close(fig)

def draw_reaction_arrow(ax, node_x, node_y, force_mag, axis_nx, axis_ny):
    if abs(force_mag) < 0.1:
        return
    arr_L = 0.5  
    sgn = np.sign(force_mag)
    dx = sgn * axis_nx
    dy = sgn * axis_ny
    start_x = node_x - arr_L * dx
    start_y = node_y - arr_L * dy
    arr_c = 'blue' if force_mag >= 0 else 'red'
    ax.arrow(
        start_x, start_y, arr_L*dx, arr_L*dy, 
        length_includes_head=True, 
        head_width=0.08, head_length=0.12, 
        fc=arr_c, ec=arr_c, lw=0.8, zorder=5
    )
    ax.text(
        start_x - 0.15*dx, start_y - 0.15*dy, 
        f"{force_mag:+.1f}", 
        color=arr_c, fontsize=7, fontname='Arial', 
        ha='center', va='center'
    )

def point_on_line(px, py, x1, y1, x2, y2, tol=1e-3):
    L2 = (x2-x1)**2 + (y2-y1)**2
    if L2 == 0:
        return False
    t = max(0, min(1, ((px-x1)*(x2-x1) + (py-y1)*(y2-y1)) / L2))
    projX = x1 + t*(x2-x1)
    projY = y1 + t*(y2-y1)
    return math.hypot(px-projX, py-projY) < tol

# =========================================================
# 1. DXF Parsing Engine (Absolute Mathematical Precision)
# =========================================================
def parse_dxf_to_data(file_bytes):
    if ezdxf is None:
        return None
    tmp_path = ""
    try:
        try:
            dxf_str = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            dxf_str = file_bytes.decode('cp1252', errors='ignore')
            
        with tempfile.NamedTemporaryFile(delete=False, suffix=".dxf", mode='w', encoding='utf-8') as tmp:
            tmp.write(dxf_str)
            tmp_path = tmp.name
            
        doc = ezdxf.readfile(tmp_path)
        msp = doc.modelspace()
        
        raw_frames = []
        raw_struts = []
        raw_supports = []
        
        def match_layer(layer, target):
            l_clean = layer.lower().replace(" ", "").replace("_", "")
            return target in l_clean
        
        for e in msp:
            lyr = e.dxf.layer
            etype = e.dxftype()
            
            if match_layer(lyr, "supp"):
                if etype in ['POINT', 'CIRCLE', 'INSERT']:
                    if etype == 'POINT':
                        raw_supports.append({'x': e.dxf.location.x, 'y': e.dxf.location.y})
                    elif etype == 'CIRCLE':
                        raw_supports.append({'x': e.dxf.center.x, 'y': e.dxf.center.y})
                    elif etype == 'INSERT':
                        raw_supports.append({'x': e.dxf.insert.x, 'y': e.dxf.insert.y})
            
            elif match_layer(lyr, "push") or match_layer(lyr, "pull"):
                entities = [e]
                if etype in ['LWPOLYLINE', 'POLYLINE']:
                    entities = list(e.virtual_entities())
                for sub_e in entities:
                    if sub_e.dxftype() == 'LINE':
                        raw_struts.append({'p1': (sub_e.dxf.start.x, sub_e.dxf.start.y), 'p2': (sub_e.dxf.end.x, sub_e.dxf.end.y)})
                    
            elif match_layer(lyr, "frame"):
                entities = [e]
                if etype in ['LWPOLYLINE', 'POLYLINE']:
                    entities = list(e.virtual_entities())
                for sub_e in entities:
                    if sub_e.dxftype() == 'LINE':
                        raw_frames.append({'type': 'line', 'p1': (sub_e.dxf.start.x, sub_e.dxf.start.y), 'p2': (sub_e.dxf.end.x, sub_e.dxf.end.y)})
                    elif sub_e.dxftype() == 'ARC':
                        c = sub_e.dxf.center
                        r = sub_e.dxf.radius
                        sa = math.radians(sub_e.dxf.start_angle)
                        ea = math.radians(sub_e.dxf.end_angle)
                        raw_frames.append({'type': 'arc', 'c': (c.x, c.y), 'r': r, 'sa': sa, 'ea': ea})

        if not raw_frames:
            return None

        if raw_supports:
            raw_supports.sort(key=lambda s: s['x'])
            dx = -raw_supports[0]['x']
            dy = -raw_supports[0]['y']
            
            for sp in raw_supports:
                sp['x'] += dx
                sp['y'] += dy
                
            for f in raw_frames:
                if f['type'] == 'line':
                    f['p1'] = (f['p1'][0] + dx, f['p1'][1] + dy)
                    f['p2'] = (f['p2'][0] + dx, f['p2'][1] + dy)
                elif f['type'] == 'arc':
                    f['c'] = (f['c'][0] + dx, f['c'][1] + dy)
                    
            for s in raw_struts:
                s['p1'] = (s['p1'][0] + dx, s['p1'][1] + dy)
                s['p2'] = (s['p2'][0] + dx, s['p2'][1] + dy)
        
        def get_min_x(f):
            if f['type'] == 'line':
                return min(f['p1'][0], f['p2'][0])
            return f['c'][0] - f['r']
            
        raw_frames.sort(key=get_min_x)

        chained_segs = []
        for f in raw_frames:
            if f['type'] == 'line':
                p_start = f['p1']
                p_end = f['p2']
                if p_start[0] > p_end[0] + 1e-5 or (abs(p_start[0] - p_end[0]) < 1e-5 and p_start[1] > p_end[1]):
                    p_start, p_end = p_end, p_start
                    
                dx_line = p_end[0]-p_start[0]
                dy_line = p_end[1]-p_start[1]
                L = math.hypot(dx_line, dy_line)
                ang = math.degrees(math.atan2(dy_line, dx_line))
                chained_segs.append({
                    'type': 'Straight Line', 
                    'Shape Type': 'Straight Line', 
                    'L': L, 
                    'start_angle': ang, 
                    'smooth': False, 
                    'is_dxf': True, 
                    'abs_p1': p_start, 
                    'abs_p2': p_end, 
                    'kappa': 0.0
                })
            elif f['type'] == 'arc':
                sa = f['sa']
                ea = f['ea']
                if ea < sa:
                    ea += 2 * math.pi
                sweep = ea - sa
                L = f['r'] * sweep
                chained_segs.append({
                    'type': 'Curve (Arc & Radius)', 
                    'Shape Type': 'Curve (Arc & Radius)', 
                    'L': L, 
                    'Radius (R) (m)': f['r'],
                    'Curvature Direction': "Arching Up ⤴ (Concave)",
                    'start_angle': math.degrees(sa + math.pi/2), 
                    'smooth': False, 
                    'is_dxf': True, 
                    'abs_c': f['c'], 
                    'abs_r': f['r'],
                    'abs_sa': sa, 
                    'abs_ea': ea, 
                    'sweep': sweep, 
                    'kappa': 1.0/f['r']
                })

        def get_closest_segment_exact(pt, segs):
            min_d = 9999.0
            best_idx = 0
            best_s = 0.0
            pt = np.array(pt)
            for idx, seg in enumerate(segs):
                L = seg.get('L', 0.0)
                if seg.get('Shape Type') == 'Straight Line' and 'abs_p1' in seg:
                    p1 = np.array(seg['abs_p1'])
                    p2 = np.array(seg['abs_p2'])
                    v = p2 - p1
                    w = pt - p1
                    c2 = np.dot(v, v)
                    ratio = np.dot(w, v) / c2 if c2 > 1e-6 else 0.0
                    ratio = max(0.0, min(1.0, ratio))
                    proj = p1 + ratio * v
                    d = np.linalg.norm(pt - proj)
                    if d < min_d:
                        min_d = d
                        best_idx = idx
                        best_s = ratio * L
                elif seg.get('Shape Type') == 'Curve (Arc & Radius)' and 'abs_c' in seg:
                    c = np.array(seg['abs_c'])
                    r = seg['abs_r']
                    v = pt - c
                    ang = math.atan2(v[1], v[0])
                    sa = seg['abs_sa']
                    sweep = seg['sweep']
                    
                    ang_norm = (ang - sa) % (2 * math.pi)
                    if ang_norm > abs(sweep):
                        ratio = 1.0 if ang_norm < math.pi else 0.0
                    else:
                        ratio = ang_norm / sweep if abs(sweep) > 1e-6 else 0.0
                    
                    ratio = max(0.0, min(1.0, ratio))
                    current_ang = sa + ratio * sweep
                    proj = c + r * np.array([math.cos(current_ang), math.sin(current_ang)])
                    d = np.linalg.norm(pt - proj)
                    if d < min_d:
                        min_d = d
                        best_idx = idx
                        best_s = ratio * L
            return min_d, best_idx, best_s

        struts_mapped = []
        for s in raw_struts:
            p1 = s['p1']
            p2 = s['p2']
            if p1[1] > p2[1]:
                top_p = p1
                bot_p = p2
            else:
                top_p = p2
                bot_p = p1
                
            d_top, b_seg, b_s = get_closest_segment_exact(top_p, chained_segs)
            struts_mapped.append({
                'seg_idx': b_seg, 
                'dist': b_s, 
                'gx': bot_p[0], 
                'gy': bot_p[1]
            })

        supps_mapped = []
        for sp in raw_supports:
            d_min, b_seg, b_s = get_closest_segment_exact((sp['x'], sp['y']), chained_segs)
            if d_min < 0.1:
                supps_mapped.append({'x': sp['x'], 'y': sp['y'], 'type': 'Hinged', 'seg_idx': b_seg, 's_dist': b_s})
            else:
                supps_mapped.append({'x': sp['x'], 'y': sp['y'], 'type': 'Hinged'})

        return {'segments': chained_segs, 'struts': struts_mapped, 'supports': supps_mapped}
        
    except Exception as e:
        st.error(f"⚠️ حدث خطأ أثناء تحليل ملف الـ DXF. تأكد من أن الملف سليم: {e}")
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except:
                pass

# =========================================================
# 2. Geometry & Mesh Generators (Absolute Eval Engine)
# =========================================================
def eval_seg_point(seg, s_val, start_data=None):
    L = seg.get('L', 0.0)
    s_val = min(max(s_val, 0.0), L)
    ratio = s_val / L if L > 1e-6 else 0.0
    
    is_dxf = seg.get('is_dxf', False)
    shape_type = seg.get('Shape Type', 'Straight Line')
    
    if is_dxf:
        if shape_type == 'Straight Line' and 'abs_p1' in seg:
            p1 = seg['abs_p1']
            p2 = seg['abs_p2']
            px = p1[0] + ratio * (p2[0] - p1[0])
            py = p1[1] + ratio * (p2[1] - p1[1])
            dx = p2[0] - p1[0]
            dy = p2[1] - p1[1]
            th = math.atan2(dy, dx)
            return px, py, th
            
        elif shape_type == 'Curve (Arc & Radius)' and 'abs_c' in seg:
            c = seg['abs_c']
            r = seg['abs_r']
            current_ang = seg['abs_sa'] + ratio * seg.get('sweep', 0)
            px = c[0] + r * math.cos(current_ang)
            py = c[1] + r * math.sin(current_ang)
            th = current_ang + math.pi/2
            return px, py, th
            
    if start_data:
        x0 = start_data.get('x0', 0)
        y0 = start_data.get('y0', 0)
        th0 = start_data.get('th0', 0)
        kappa = start_data.get('kappa', 0)
        
        if abs(kappa) < 1e-6: 
            x = x0 + s_val * math.cos(th0)
            y = y0 + s_val * math.sin(th0)
            th = th0
        else: 
            x = x0 + (math.sin(th0 + kappa * s_val) - math.sin(th0)) / kappa
            y = y0 - (math.cos(th0 + kappa * s_val) - math.cos(th0)) / kappa
            th = th0 + kappa * s_val
        return x, y, th
    
    return 0.0, 0.0, 0.0

def get_approx_xy(segs, s_idx, s_val):
    if s_idx < 0 or s_idx >= len(segs):
        return 0.0, 0.0
    seg = segs[s_idx]
    if seg.get('is_dxf'):
        px, py, _ = eval_seg_point(seg, s_val)
        return px, py
        
    curr_x = 0.0
    curr_y = 0.0
    curr_th = 0.0
    
    for i in range(s_idx + 1):
        sg = segs[i]
        if i == 0 or not sg.get('smooth', True):
            curr_th = math.radians(sg.get('start_angle', 0.0))
        L = sg.get('L', 0.0)
        kappa = sg.get('kappa', 0.0)
        
        if i == s_idx:
            if abs(kappa) < 1e-6:
                return curr_x + s_val * math.cos(curr_th), curr_y + s_val * math.sin(curr_th)
            else:
                x_val = curr_x + (math.sin(curr_th + kappa*s_val) - math.sin(curr_th))/kappa
                y_val = curr_y - (math.cos(curr_th + kappa*s_val) - math.cos(curr_th))/kappa
                return x_val, y_val
                
        if abs(kappa) < 1e-6:
            curr_x += L * math.cos(curr_th)
            curr_y += L * math.sin(curr_th)
        else:
            curr_x += (math.sin(curr_th + kappa*L) - math.sin(curr_th))/kappa
            curr_y -= (math.cos(curr_th + kappa*L) - math.cos(curr_th))/kappa
            curr_th += kappa * L
            
    return curr_x, curr_y

def build_chain_mesh(segments, seg_sections, loads, struts, base_sec, supports, corner_sup, mesh_size=0.25):
    nodes = []
    elements = []
    nodal_loads = []
    
    node_tol = 1e-4 
    
    def get_or_add_node(x, y):
        for i, n in enumerate(nodes):
            if abs(n[0] - x) < node_tol and abs(n[1] - y) < node_tol:
                return i
        nodes.append([x, y])
        return len(nodes) - 1

    seg_start_data = [] 
    curr_x = 0.0
    curr_y = 0.0
    curr_th = np.radians(corner_sup.get('angle', 0.0)) 
    
    key_nodes = set()
    
    for i, seg in enumerate(segments):
        if not seg.get('is_dxf'):
            if i == 0 or not seg.get('smooth', True):
                curr_th = np.radians(seg.get('start_angle', 0.0))
            
        L = seg.get('L', 0.0)
        kappa = seg.get('kappa', 0.0)
        seg_start_data.append({'x0': curr_x, 'y0': curr_y, 'th0': curr_th, 'kappa': kappa})
        
        key_s_vals = [0.0, L]
        for st_item in struts:
            if st_item.get('seg_idx') == i:
                key_s_vals.append(st_item['s_dist'])
        for ld in loads:
            if ld.get('seg_idx') == i:
                key_s_vals.append(ld['start'])
                key_s_vals.append(ld['end'])
        for sp in supports:
            if sp.get('seg_idx') == i:
                key_s_vals.append(sp['s_dist'])
            
        keys = list(key_s_vals)
        num_sub = max(1, int(np.ceil(L / mesh_size)))
        for p in np.linspace(0, L, num_sub+1):
            keys.append(p)
            
        keys = sorted(list(set([min(max(round(k, 5), 0.0), round(L, 5)) for k in keys])))
        
        node_indices = []
        for s_val in keys:
            px, py, _ = eval_seg_point(seg, s_val, seg_start_data[i])
            nid = get_or_add_node(px, py)
            node_indices.append(nid)
            if any(abs(s_val - kv) < 1e-4 for kv in key_s_vals):
                key_nodes.add(nid)
            
        sec_props = seg_sections[i]
        
        for j in range(len(keys)-1):
            n1 = node_indices[j]
            n2 = node_indices[j+1]
            if n1 == n2:
                continue 
                
            s_mid = (keys[j] + keys[j+1]) / 2.0
            _, _, th_mid = eval_seg_point(seg, s_mid, seg_start_data[i])
            c_t = np.cos(th_mid)
            s_t = np.sin(th_mid)
            
            p_x1 = 0.0
            p_y1 = 0.0
            p_x2 = 0.0
            p_y2 = 0.0
            
            for ld in loads:
                if ld.get('seg_idx') == i and ld.get('type') != 'Point Load':
                    if ld['start'] - 1e-4 <= s_mid <= ld['end'] + 1e-4:
                        L_ld = max(ld['end'] - ld['start'], 1e-5)
                        wa = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j] - ld['start']) / L_ld
                        wb = ld['w1'] + (ld['w2'] - ld['w1']) * (keys[j+1] - ld['start']) / L_ld
                        
                        dir_str = ld.get('dir', '')
                        if 'Global Z' in dir_str or 'Global Y' in dir_str:
                            p_x1 += wa * s_t
                            p_y1 += wa * c_t
                            p_x2 += wb * s_t
                            p_y2 += wb * c_t
                        elif 'Global X' in dir_str:
                            p_x1 += wa * c_t
                            p_y1 -= wa * s_t
                            p_x2 += wb * c_t
                            p_y2 -= wb * s_t
                        else:
                            p_x1 += 0.0
                            p_y1 += wa
                            p_x2 += 0.0
                            p_y2 += wb
                            
            elements.append({
                'type': 'frame', 
                'group': 'segment', 
                'sec': sec_props['name'],
                'n1': n1, 
                'n2': n2, 
                'px1': p_x1, 
                'py1': p_y1, 
                'px2': p_x2, 
                'py2': p_y2,
                'E': sec_props['E'] * 10000.0, 
                'A': sec_props['A'], 
                'I': sec_props['I'] / 100000000.0,
                'seg_idx': i, 
                's_start': keys[j], 
                's_end': keys[j+1], 
                'L': keys[j+1] - keys[j],
                'th_mid': th_mid
            })
            
        for ld in loads:
            if ld.get('seg_idx') == i and ld.get('type') == 'Point Load':
                try:
                    idx = keys.index(round(ld['start'], 5))
                    nid = node_indices[idx]
                    dir_str = ld.get('dir', '')
                    if 'Global Z' in dir_str or 'Global Y' in dir_str:
                        nodal_loads.append({'node': nid, 'Fx': 0.0, 'Fy': ld['w1']})
                    elif 'Global X' in dir_str:
                        nodal_loads.append({'node': nid, 'Fx': ld['w1'], 'Fy': 0.0})
                    else: 
                        _, _, th_pt = eval_seg_point(seg, ld['start'], seg_start_data[i])
                        c_pt = np.cos(th_pt)
                        s_pt = np.sin(th_pt)
                        nodal_loads.append({'node': nid, 'Fx': -ld['w1']*s_pt, 'Fy': ld['w1']*c_pt})
                except ValueError:
                    pass
                
        curr_x, curr_y, curr_th = eval_seg_point(seg, L, seg_start_data[i])

    for st_idx, st_item in enumerate(struts):
        seg_idx = st_item.get('seg_idx', 0)
        dist = round(st_item.get('s_dist', 0.0), 5)
        gx = st_item.get('gx', 0.0)
        gy = st_item.get('gy', 0.0)
        
        nx, ny, _ = eval_seg_point(segments[seg_idx], dist, seg_start_data[seg_idx])
        
        top_node = get_or_add_node(nx, ny)
        bot_node = get_or_add_node(gx, gy)
        
        elements.append({
            'type': 'truss', 
            'group': 'strut', 
            'sec': st_item.get('sec', 'Unknown'),
            'n1': bot_node, 
            'n2': top_node, 
            'strut_idx': st_idx,
            'E': 21000000.0, 
            'A': 0.001
        })

    supports_list = []
    for sup in supports:
        if 'seg_idx' in sup:
            s_dist = round(sup['s_dist'], 5)
            nx, ny, _ = eval_seg_point(segments[sup['seg_idx']], s_dist, seg_start_data[sup['seg_idx']])
            nid = get_or_add_node(nx, ny)
        else:
            nid = get_or_add_node(sup.get('x', 0.0), sup.get('y', 0.0))
        supports_list.append({'node': nid, 'type': sup.get('type', 'Hinged'), 'angle': 0.0})
        
    display_nodes = set([s['node'] for s in supports_list])
    display_nodes.update(key_nodes) 

    return nodes, elements, nodal_loads, display_nodes, supports_list, seg_start_data

# =========================================================
# V4 AUTO-MESH ENGINE (For Interactive CAD Output)
# =========================================================
def build_cad_mesh(cad_payload):
    cad_lines = cad_payload.get('cad_lines', [])
    supports = cad_payload.get('supports', [])
    struts = cad_payload.get('struts', [])
    loads = cad_payload.get('loads', [])
    overrides = cad_payload.get('overrides', {})
    combo = cad_payload.get('combo', {'D': 1.0, 'L': 0.0, 'W': 0.0})
    
    nodes = []
    elements = []
    nodal_loads = []
    supports_list = []
    display_nodes = set()
    
    def get_or_add_node(x, y):
        for i, n in enumerate(nodes):
            if abs(n[0] - x) < 1e-3 and abs(n[1] - y) < 1e-3:
                return i
        nodes.append([x, y])
        return len(nodes) - 1

    cut_points = []
    for sp in supports:
        cut_points.append((sp['x']/100, -sp['y']/100))
        
    for st in struts:
        cut_points.append((st['gx']/100, -st['gy']/100))
        cut_points.append((st['top_x']/100, -st['top_y']/100))

    meshed_frames = []
    for line in cad_lines:
        x1 = line['x1']/100
        y1 = -line['y1']/100
        x2 = line['x2']/100
        y2 = -line['y2']/100
        
        pts_on_line = [(x1, y1), (x2, y2)]
        for cp in cut_points:
            if point_on_line(cp[0], cp[1], x1, y1, x2, y2):
                pts_on_line.append((cp[0], cp[1]))
                
        pts_on_line.sort(key=lambda p: math.hypot(p[0]-x1, p[1]-y1))
        
        unique_pts = []
        for p in pts_on_line:
            if not unique_pts or math.hypot(p[0]-unique_pts[-1][0], p[1]-unique_pts[-1][1]) > 1e-4:
                unique_pts.append(p)
                
        for i in range(len(unique_pts)-1):
            meshed_frames.append({
                'id': line['id'], 
                'curve_id': line.get('curve_id'),
                'x1': unique_pts[i][0], 
                'y1': unique_pts[i][1],
                'x2': unique_pts[i+1][0], 
                'y2': unique_pts[i+1][1]
            })

    seg_sections = []
    for m_line in meshed_frames:
        n1 = get_or_add_node(m_line['x1'], m_line['y1'])
        n2 = get_or_add_node(m_line['x2'], m_line['y2'])
        
        if n1 != n2:
            sec_name = overrides.get(m_line['id'], {}).get('sec', "Soldier U100")
            sec = SECTIONS_DB.get(sec_name, {"E": 2100.0, "A": 0.0034, "I": 0.00000412, "Mall": 13.1, "Qall": 100.8})
            seg_sections.append({'name': sec_name, 'Mall': sec.get('Mall', 13.1), 'Qall': sec.get('Qall', 100.8)})
            L = np.hypot(nodes[n2][0] - nodes[n1][0], nodes[n2][1] - nodes[n1][1])
            
            elements.append({
                'id': m_line['id'], 
                'curve_id': m_line.get('curve_id'), 
                'type': 'frame', 
                'group': 'segment',
                'n1': n1, 
                'n2': n2, 
                'E': sec.get('E', 2100.0)*10000, 
                'A': sec.get('A', 0.0034), 
                'I': sec.get('I', 0.0000041), 
                'c': (nodes[n2][0]-nodes[n1][0])/L, 
                's': (nodes[n2][1]-nodes[n1][1])/L, 
                'L': L, 
                'px1': 0.0, 
                'py1': 0.0, 
                'px2': 0.0, 
                'py2': 0.0
            })

    for st in struts:
        n1 = get_or_add_node(st['gx']/100, -st['gy']/100)
        n2 = get_or_add_node(st['top_x']/100, -st['top_y']/100)
        sec_name = overrides.get(st['id'], {}).get('sec', "PPH 353")
        
        elements.append({
            'id': st['id'], 
            'type': 'truss', 
            'group': 'strut', 
            'sec': sec_name, 
            'n1': n1, 
            'n2': n2, 
            'E': 21000000.0, 
            'A': 0.002
        })

    for sp in supports:
        node_idx = get_or_add_node(sp['x']/100, -sp['y']/100)
        sup_type = overrides.get(sp['id'], {}).get('type', 'Hinged')
        sup_ang = float(overrides.get(sp['id'], {}).get('angle', 0.0))
        
        supports_list.append({
            'node': node_idx, 
            'type': sup_type, 
            'angle': sup_ang
        })
        display_nodes.add(node_idx)

    for ld in loads:
        factored_w1 = float(ld['w1']) * combo.get(ld.get('case', 'D'), 1.0)
        factored_w2 = float(ld.get('w2', ld['w1'])) * combo.get(ld.get('case', 'D'), 1.0)
        
        if factored_w1 == 0 and factored_w2 == 0:
            continue
        
        for el in elements:
            if el['type'] == 'frame' and (ld['target'] == 'all' or el['id'] in ld['target'] or el.get('curve_id') in ld['target']):
                dir_str = ld.get('dir', 'Vertical')
                if 'Vertical' in dir_str:
                    el['py1'] -= factored_w1
                    el['py2'] -= factored_w2 
                elif 'Horizontal' in dir_str:
                    el['px1'] += factored_w1
                    el['px2'] += factored_w2
                else:
                    el['px1'] -= el['s']*factored_w1
                    el['py1'] += el['c']*factored_w1
                    el['px2'] -= el['s']*factored_w2
                    el['py2'] += el['c']*factored_w2

    return nodes, elements, nodal_loads, display_nodes, supports_list, seg_sections

# =========================================================
# 3. Advanced FEA Solver
# =========================================================
def solve_fea_engine(nodes, elements, nodal_loads, supports_list):
    NDOF = len(nodes) * 3
    K = np.zeros((NDOF, NDOF))
    F = np.zeros(NDOF)
    
    for el in elements:
        n1 = el['n1']
        n2 = el['n2']
        x1 = nodes[n1][0]
        y1 = nodes[n1][1]
        x2 = nodes[n2][0]
        y2 = nodes[n2][1]
        
        L = np.hypot(x2 - x1, y2 - y1)
        
        if L < 1e-5: 
            el['c'] = 1
            el['s'] = 0
            el['L'] = 1e-5
            el['internal'] = {'N': [0,0], 'V': [0,0], 'M': [0,0], 'x': [0, 1e-5], 'v_rel': [0,0]}
            continue
            
        c = (x2 - x1) / L
        s = (y2 - y1) / L
        el['L'] = L
        el['c'] = c
        el['s'] = s
        
        E = el['E']
        A = el['A']
        I = el.get('I', 0.00005)
        
        T = np.array([
            [c, s, 0, 0, 0, 0], 
            [-s, c, 0, 0, 0, 0], 
            [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], 
            [0, 0, 0, -s, c, 0], 
            [0, 0, 0, 0, 0, 1]
        ])
        
        if el['type'] == 'truss':
            k_loc = np.zeros((6, 6))
            k_loc[0,0] = E*A/L
            k_loc[3,3] = E*A/L
            k_loc[0,3] = -E*A/L
            k_loc[3,0] = -E*A/L
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1 = el.get('px1',0)
            py1 = el.get('py1',0)
            px2 = el.get('px2',0)
            py2 = el.get('py2',0)
            
            f_loc = np.array([
                (2*px1 + px2)*L/6.0, 
                (7*py1 + 3*py2)*L/20.0, 
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0, 
                (3*py1 + 7*py2)*L/20.0, 
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            f_glob = T.T @ f_loc
            dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
            for r in range(6): 
                F[dof[r]] += f_glob[r]
            
        k_glob = T.T @ k_loc @ T
        dof = [3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]
        for r in range(6):
            for col in range(6): 
                K[dof[r], dof[col]] += k_glob[r, col]
                
    for nl in nodal_loads:
        F[3*nl['node']] += nl['Fx']
        F[3*nl['node']+1] += nl['Fy']
            
    K_orig = K.copy()
    fixed_dofs = []
    K_pen = 1e12
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        a = sup.get('angle', 0.0)
        
        if t == 'Fixed':
            fixed_dofs.extend([3*n, 3*n+1, 3*n+2])
        elif t == 'Hinged':
            fixed_dofs.extend([3*n, 3*n+1])
        elif t == 'Roller':
            rad = np.radians(a)
            nx = -np.sin(rad)
            ny = np.cos(rad) 
            K[3*n, 3*n] += K_pen*nx**2
            K[3*n+1, 3*n+1] += K_pen*ny**2
            K[3*n, 3*n+1] += K_pen*nx*ny
            K[3*n+1, 3*n] += K_pen*nx*ny

    free_dof = [i for i in range(NDOF) if i not in fixed_dofs]
    K_ff = K[np.ix_(free_dof, free_dof)]
    F_f = F[free_dof]
    
    U = np.zeros(NDOF)
    try: 
        U[free_dof] = np.linalg.solve(K_ff, F_f)
    except np.linalg.LinAlgError: 
        U[free_dof] = np.linalg.lstsq(K_ff, F_f, rcond=None)[0]
    
    R_reactions = K_orig @ U - F 
    
    for el in elements:
        if el.get('L', 0) < 1e-5:
            continue
            
        n1 = el['n1']
        n2 = el['n2']
        c = el['c']
        s = el['s']
        L = el['L']
        E = el['E']
        A = el['A']
        I = el.get('I', 0.00005)
        
        u_glob = U[[3*n1, 3*n1+1, 3*n1+2, 3*n2, 3*n2+1, 3*n2+2]]
        T = np.array([
            [c, s, 0, 0, 0, 0], 
            [-s, c, 0, 0, 0, 0], 
            [0, 0, 1, 0, 0, 0],
            [0, 0, 0, c, s, 0], 
            [0, 0, 0, -s, c, 0], 
            [0, 0, 0, 0, 0, 1]
        ])
        u_loc = T @ u_glob
        el['internal'] = {'u_loc': u_loc}
        
        if el['type'] == 'truss':
            N_val = (E * A / L) * (u_loc[3] - u_loc[0])
            xs = np.linspace(0, L, 51)
            el['internal'].update({
                'N': np.full_like(xs, N_val),
                'V': np.zeros_like(xs),
                'M': np.zeros_like(xs),
                'x': xs,
                'v_rel': np.zeros_like(xs)
            })
        else:
            k_loc = np.array([
                [E*A/L, 0, 0, -E*A/L, 0, 0], 
                [0, 12*E*I/L**3, 6*E*I/L**2, 0, -12*E*I/L**3, 6*E*I/L**2],
                [0, 6*E*I/L**2, 4*E*I/L, 0, -6*E*I/L**2, 2*E*I/L], 
                [-E*A/L, 0, 0, E*A/L, 0, 0],
                [0, -12*E*I/L**3, -6*E*I/L**2, 0, 12*E*I/L**3, -6*E*I/L**2], 
                [0, 6*E*I/L**2, 2*E*I/L, 0, -6*E*I/L**2, 4*E*I/L]
            ])
            
            px1 = el.get('px1',0)
            py1 = el.get('py1',0)
            px2 = el.get('px2',0)
            py2 = el.get('py2',0)
            
            f_loc = np.array([
                (2*px1 + px2)*L/6.0, 
                (7*py1 + 3*py2)*L/20.0, 
                (3*py1 + 2*py2)*L**2/60.0,
                (px1 + 2*px2)*L/6.0, 
                (3*py1 + 7*py2)*L/20.0, 
                -(2*py1 + 3*py2)*L**2/60.0
            ])
            
            f_end = k_loc @ u_loc - f_loc
            
            xs = np.linspace(0, L, 51) 
            N_arr = np.zeros_like(xs)
            V_arr = np.zeros_like(xs)
            M_arr = np.zeros_like(xs)
            v_rel_arr = np.zeros_like(xs)
            
            v1 = u_loc[1]
            theta1 = u_loc[2]
            v2 = u_loc[4]
            theta2 = u_loc[5]
            
            for i, x in enumerate(xs):
                N_arr[i] = -f_end[0] - (px1*x + (px2-px1)*x**2/(2*L))
                V_arr[i] = f_end[1] + (py1*x + (py2-py1)*x**2/(2*L))
                M_arr[i] = -f_end[2] + f_end[1]*x + py1*x**2/2.0 + (py2-py1)*x**3/(6*L)
                
                xi = x / L
                v_shape = v1*(1-3*xi**2+2*xi**3) + theta1*(x*(1-xi)**2) + v2*(3*xi**2-2*xi**3) + theta2*(x*(xi**2-xi))
                v_chord = v1 + xi * (v2 - v1) 
                v_rel_arr[i] = v_shape - v_chord 
                
            el['internal'].update({'N': N_arr, 'V': V_arr, 'M': M_arr, 'x': xs, 'v_rel': v_rel_arr})
            
    return U, R_reactions

# =========================================================
# 4. Plotting Engine (Live Preview & SAP2000 Colors)
# =========================================================
def draw_base_geometry(ax, nodes, elements, supports_list, seg_sections=None, segments=None, seg_starts=None):
    for el in elements:
        if el['type'] not in ['frame', 'truss']:
            continue
            
        n1 = nodes[el['n1']]
        n2 = nodes[el['n2']]
        
        if el['type'] == 'truss':
            ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='red', linestyle='-', linewidth=0.8, zorder=1)
        else:
            if el.get('group') == 'base' and el.get('sec') == "None (Direct to Ground)":
                continue
                
            if el.get('group') == 'segment' and segments and seg_starts:
                s_idx = el['seg_idx']
                seg = segments[s_idx]
                s_data = seg_starts[s_idx]
                curve_x = []
                curve_y = []
                s_start = el.get('s_start', 0.0)
                s_end = el.get('s_end', el.get('L', 0.0))
                
                for p in np.linspace(s_start, s_end, 10):
                    cx, cy, _ = eval_seg_point(seg, p, s_data)
                    curve_x.append(cx)
                    curve_y.append(cy)
                ax.plot(curve_x, curve_y, color='royalblue', linestyle='-', linewidth=1.5, zorder=1)
            else:
                ax.plot([n1[0], n2[0]], [n1[1], n2[1]], color='royalblue', linestyle='-', linewidth=1.5, zorder=1)
            
    for i, sup in enumerate(supports_list):
        n = sup['node']
        x = nodes[n][0]
        y = nodes[n][1]
        t = sup['type']
        
        ax.text(x, y - 0.4, f"J{i+1}", color='green', fontsize=7, ha='center', fontname='Arial')
        
        if t == 'Fixed':
            ax.plot(x, y, marker='s', markerfacecolor='none', markeredgecolor='limegreen', markersize=3, zorder=5)
            ax.plot([x - 0.1, x + 0.1], [y - 0.1, y + 0.1], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Hinged':
            h = 0.15
            w = 0.12
            p1 = (x, y)
            p2 = (x + w, y - h)
            p3 = (x - w, y - h)
            ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            ax.plot([x - w - 0.05, x + w + 0.05], [y - h, y - h], color='limegreen', lw=1.0, zorder=4)
        elif t == 'Roller':
            h = 0.15
            w = 0.12
            r = 0.04
            p1 = (x, y)
            p2 = (x + w, y - h)
            p3 = (x - w, y - h)
            ax.add_patch(Polygon([p1, p2, p3], facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            ax.add_patch(plt.Circle((x, y - h - r), r, facecolor='none', edgecolor='limegreen', lw=1.0, zorder=5))
            ax.plot([x - w - 0.05, x + w + 0.05], [y - h - 2*r, y - h - 2*r], color='limegreen', lw=1.0, zorder=4)

    if seg_sections and segments and seg_starts:
        for el in elements:
            if el['type'] == 'truss':
                n1 = nodes[el['n1']]
                n2 = nodes[el['n2']]
                mid_x = (n1[0]+n2[0])/2
                mid_y = (n1[1]+n2[1])/2
                dx = n2[0]-n1[0]
                dy = n2[1]-n1[1]
                rot = np.degrees(math.atan2(dy, dx))
                
                if rot > 90:
                    rot -= 180
                elif rot < -90:
                    rot += 180
                    
                L_hyp = np.hypot(dx, dy)
                if L_hyp > 1e-4:
                    nx_s = -dy/L_hyp
                    ny_s = dx/L_hyp
                    st_id = el.get('strut_idx', 0) + 1
                    label = f"P{st_id}: {get_short_name(el.get('sec', ''))}"
                    ax.text(
                        mid_x + nx_s*0.1, mid_y + ny_s*0.1, label, 
                        color='dimgray', fontsize=6, rotation=rot, 
                        ha='center', va='center', fontname='Arial'
                    )
        
        for i, seg in enumerate(segments):
            s_data = seg_starts[i]
            mx, my, mth = eval_seg_point(seg, seg.get('L', 0)/2, s_data)
            rot_deg = math.degrees(mth)
            
            if rot_deg > 90:
                rot_deg -= 180
            elif rot_deg < -90:
                rot_deg += 180
                
            sec_name = seg_sections[i]['name']
            label = f"S{i+1}: {get_short_name(sec_name)}"
            ax.text(
                mx - math.sin(mth)*0.1, my + math.cos(mth)*0.1, label, 
                color='dimgray', fontsize=6, ha='center', va='center', 
                rotation=rot_deg, fontname='Arial'
            )

def draw_loads_and_geometry(ax, nodes, elements, supports_list, seg_sections, loads, segments=None, seg_starts=None):
    draw_base_geometry(ax, nodes, elements, supports_list, seg_sections, segments, seg_starts)

    scale_ld = 0.05
    for ld in loads:
        if segments and seg_starts:
            i = ld.get('seg_idx', 0)
            s_data = seg_starts[i]
            w1 = ld.get('w1', 0)
            w2 = ld.get('w2', 0)
            
            num_pts = max(10, int((ld.get('end', 0) - ld.get('start', 0)) / 0.1))
            s_vals = np.linspace(ld.get('start', 0), ld.get('end', 0), num_pts)
            poly_pts = []
            top_pts = []
            
            for sv in s_vals:
                px, py, th = eval_seg_point(segments[i], sv, s_data)
                w_curr = w1 + (w2 - w1) * (sv - ld.get('start', 0)) / max(1e-5, (ld.get('end', 0) - ld.get('start', 0)))
                w_val = w_curr * scale_ld
                poly_pts.append((px, py))
                
                dir_str = ld.get('dir', '')
                if 'Global Z' in dir_str or 'Global Y' in dir_str:
                    f_vx = 0.0
                    f_vy = w_val
                elif 'Global X' in dir_str:
                    f_vx = w_val
                    f_vy = 0.0
                else:
                    c = math.cos(th)
                    s = math.sin(th)
                    f_vx = -s * w_val
                    f_vy = c * w_val
                    
                top_pts.append((px - f_vx, py - f_vy))
                    
            poly_pts.extend(top_pts[::-1])
            if len(poly_pts) > 2:
                ax.add_patch(Polygon(poly_pts, facecolor='none', edgecolor='blue', lw=0.8, zorder=2))

def draw_live_preview(nodes, elements, supports_list, seg_sections, loads, segments=None, seg_starts=None):
    apply_plot_styles()
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_aspect('equal', adjustable='datalim')
    ax.axis('off')
    draw_loads_and_geometry(ax, nodes, elements, supports_list, seg_sections, loads, segments, seg_starts)
    return safe_render_fig(fig)

def plot_sap2000_diagrams(nodes, elements, R_reactions, scales, display_nodes, supports_list, seg_sections, loads, segments=None, seg_starts=None):
    apply_plot_styles()
    figs_dict = {}
    
    fig_ld, ax_ld = plt.subplots(figsize=(6, 5))
    ax_ld.set_aspect('equal', adjustable='datalim')
    ax_ld.axis('off')
    draw_loads_and_geometry(ax_ld, nodes, elements, supports_list, seg_sections, loads, segments, seg_starts)
    figs_dict['Load'] = safe_render_fig(fig_ld)
    
    fig_r, ax_r = plt.subplots(figsize=(6, 5))
    ax_r.set_aspect('equal', adjustable='datalim')
    ax_r.axis('off')
    draw_base_geometry(ax_r, nodes, elements, supports_list, seg_sections, segments, seg_starts)
    
    for sup in supports_list:
        n = sup['node']
        t = sup['type']
        Rx = R_reactions[3*n]
        Ry = R_reactions[3*n+1]
        x = nodes[n][0]
        y = nodes[n][1]
        
        if t == 'Roller':
            draw_reaction_arrow(ax_r, x, y, Ry, 0, 1)
        else:
            draw_reaction_arrow(ax_r, x, y, Rx, 1, 0)
            draw_reaction_arrow(ax_r, x, y, Ry, 0, 1)
            
    figs_dict['React'] = safe_render_fig(fig_r)
    
    def create_force_plot(val_key, scale, c_pos, c_neg):
        fig_f, ax_f = plt.subplots(figsize=(6, 5))
        ax_f.set_aspect('equal', adjustable='datalim')
        ax_f.axis('off')
        draw_base_geometry(ax_f, nodes, elements, supports_list, seg_sections, segments, seg_starts)
        
        global_texts = []
        def is_far(tx, ty):
            for (px, py) in global_texts:
                if math.hypot(tx-px, ty-py) < 0.35:
                    return False
            return True

        global_max = 0.0
        for el in elements:
            vals = el.get('internal', {}).get(val_key, [0])
            global_max = max(global_max, np.max(np.abs(vals)))

        for el in elements:
            n1 = el['n1']
            n2 = el['n2']
            x1 = nodes[n1][0]
            y1 = nodes[n1][1]
            x2 = nodes[n2][0]
            y2 = nodes[n2][1]
            c = el.get('c', 1)
            s = el.get('s', 0)
            L = el.get('L', 0)
            
            xs = el.get('internal', {}).get('x', [])
            vals = el.get('internal', {}).get(val_key, [])
            
            if len(vals) == 0 or np.all(np.abs(vals) < 1e-4):
                continue
            
            plot_vals = -vals if val_key != 'N' else vals
            px = x1 + c * xs - s * plot_vals * scale
            py = y1 + s * xs + c * plot_vals * scale
            
            for k in range(len(px)-1):
                color = c_pos if vals[k] >= 0 else c_neg
                ax_f.plot([px[k], px[k+1]], [py[k], py[k+1]], color=color, lw=0.8)
                
            ax_f.plot([x1, px[0]], [y1, py[0]], color=c_pos if vals[0]>=0 else c_neg, lw=0.8)
            ax_f.plot([x2, px[-1]], [y2, py[-1]], color=c_pos if vals[-1]>=0 else c_neg, lw=0.8)

            def plot_val(idx):
                v = vals[idx]
                if abs(v) < 0.1:
                    return
                tx = px[idx]
                ty = py[idx]
                sgn = 1 if plot_vals[idx] >= 0 else -1
                tx += -s * sgn * 0.15
                ty += c * sgn * 0.15
                v_color = c_pos if v >= 0 else c_neg
                if is_far(tx, ty):
                    ax_f.text(tx, ty, f"{v:+.1f}", fontsize=6, color=v_color, ha='center', va='center', fontname='Arial')
                    global_texts.append((tx, ty))

            if len(vals) > 0:
                plot_val(len(vals)//2)
                
        return safe_render_fig(fig_f)

    figs_dict['N'] = create_force_plot('N', scales['N'], 'blue', 'red')
    figs_dict['V'] = create_force_plot('V', scales['V'], 'blue', 'red')
    figs_dict['M'] = create_force_plot('M', scales['M'], 'blue', 'red')
    
    fig_d, ax_d = plt.subplots(figsize=(6, 5))
    ax_d.set_aspect('equal', adjustable='datalim')
    ax_d.axis('off')
    draw_base_geometry(ax_d, nodes, elements, supports_list, seg_sections, segments, seg_starts)
    
    max_def = 0.0
    for el in elements:
        if el['type'] == 'frame':
            xs = el.get('internal', {}).get('x', [])
            v_rel = el.get('internal', {}).get('v_rel', np.zeros_like(xs)) * 20.0 
            if len(xs) == 0:
                continue
            
            x1 = nodes[el['n1']][0]
            y1 = nodes[el['n1']][1]
            c = el.get('c', 1)
            s = el.get('s', 0)
            px = x1 + c * xs - s * v_rel
            py = y1 + s * xs + c * v_rel
            ax_d.plot(px, py, color='red', linestyle='--', linewidth=1.2, alpha=0.8)
            
            max_def_local = np.max(np.abs(el.get('internal', {}).get('v_rel', [0])))
            max_def = max(max_def, max_def_local)
            
    if max_def > 0:
        ax_d.text(
            nodes[0][0], nodes[0][1]+1.0, 
            f"Max Deflection = {max_def*1000:.2f} mm", 
            color='red', fontsize=10, fontweight='bold'
        )
        
    figs_dict['D'] = safe_render_fig(fig_d)

    return figs_dict

# =========================================================
# 5. Word Report Generator
# =========================================================
def generate_chain_report(sys_data):
    if os.path.exists("Acrow_Template.docx"):
        doc = Document("Acrow_Template.docx")
        doc.add_page_break()
    else:
        doc = Document()
        
    def force_ltr_left(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '0')
        pPr.append(bidi)
        
    def add_line(text, bold=False):
        p = doc.add_paragraph()
        force_ltr_left(p)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = bold
        r.font.rtl = False
        
    p_title = doc.add_paragraph()
    force_ltr_left(p_title)
    run_title = p_title.add_run("CALCULATION SHEET FOR ADVANCED SHAPES")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.rtl = False
    
    add_line("="*50, bold=True)
    
    add_line(f"1. Safety Checks:", bold=True)
    for df_row in sys_data['safety_df']:
        add_line(f"- {df_row['Component']} ({df_row['Force Type']}): {df_row['Actual']} vs {df_row['Allowable']} => {df_row['Status']}")
    
    doc.add_page_break()
    add_line("2. Analysis Diagrams:", bold=True)
    doc.add_paragraph()
    
    def add_diagram_pair(doc, img1_bytes, title1, img2_bytes, title2):
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        p1 = table.rows[0].cells[0].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run().add_picture(io.BytesIO(img1_bytes), width=Cm(8.0))
        
        p2 = table.rows[0].cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run().add_picture(io.BytesIO(img2_bytes), width=Cm(8.0))
        
        p3 = table.rows[1].cells[0].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(title1)
        r3.font.name = 'Arial'
        r3.font.size = Pt(10)
        r3.bold = True
        
        p4 = table.rows[1].cells[1].paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(title2)
        r4.font.name = 'Arial'
        r4.font.size = Pt(10)
        r4.bold = True

    bufs = sys_data['img_bufs']
    add_diagram_pair(doc, bufs['Load'], "Assigned Load Diagram", bufs['React'], "Reactions Diagram (kN)")
    doc.add_paragraph()
    add_diagram_pair(doc, bufs['N'], "Axial Force Diagram (kN)", bufs['V'], "Shear Force Diagram (kN)")
    doc.add_paragraph()
    add_diagram_pair(doc, bufs['M'], "Bending Moment Diagram (kN.m)", bufs['D'], "Deflection Shape")
    
    out = io.BytesIO()
    doc.save(out)
    return out

def reset_adv_state():
    if 'adv_solved' in st.session_state:
        st.session_state.adv_solved = False

# =========================================================
# 6. Main Streamlit UI (The Hybrid Builder)
# =========================================================
def render_advanced_shape_module():
    st.markdown("## 🎢 The Chain Builder (Interactive CAD & Classic Mode)")
    
    input_mode = st.radio(
        "Select Input Method:", 
        ["🖱️ Interactive 2D CAD (New!)", "📝 Classic Manual Input (V3)"], 
        horizontal=True, 
        on_change=reset_adv_state
    )

    if input_mode == "🖱️ Interactive 2D CAD (New!)":
        st.info("💡 **Interactive Mode:** Draw, snap, and assign loads directly on the black screen below. Click 'Send to Python & Solve' when ready.")
        
        parent_dir = os.path.dirname(os.path.abspath(__file__))
        build_dir = os.path.join(parent_dir, "cad_component")
        
        if not os.path.exists(build_dir) or not os.path.exists(os.path.join(build_dir, "index.html")):
            st.error(f"❌ Could not find 'cad_component/index.html'. Please create this folder in the root directory and add the HTML code.")
            return

        try:
            cad_ui = components.declare_component("cad_ui", path=build_dir)
            cad_result = cad_ui(key="interactive_cad_canvas")
        except Exception as e:
            st.error(f"Error loading custom component: {e}")
            return

        if cad_result and cad_result.get("status") != "ready":
            st.success("✅ Geometry & Loads Received Successfully!")
            
            with st.spinner("Generating Auto-Mesh & Solving Matrix..."):
                from main import build_cad_mesh 
                nodes, elements, nodal_loads, display_nodes, supports_list, seg_sections = build_cad_mesh(cad_result)
                U, R = solve_fea_engine(nodes, elements, nodal_loads, supports_list)
                
                st.session_state.adv_fea_data = {
                    'U': U, 
                    'R': R, 
                    'nodes': nodes, 
                    'elements': elements, 
                    'display_nodes': display_nodes,
                    'supports_list': supports_list, 
                    'seg_sections': seg_sections,
                    'loads_data': cad_result.get('loads', [])
                }
                st.session_state.adv_solved = True
                
    else:
        # =======================================================
        # CLASSIC MODE
        # =======================================================
        view_plane = st.radio(
            "📐 Structural Analysis Plane / System Projection", 
            ["Section View", "Plan View"], 
            horizontal=True, 
            on_change=reset_adv_state
        )
        
        c_upload, c_mesh = st.columns([2, 1])
        with c_upload:
            uploaded_dxf = st.file_uploader("📥 Upload DXF File (.dxf)", type=['dxf'], key="dxf_uploader")
            
        with c_mesh:
            st.write("")
            st.write("")
            auto_mesh_size = st.number_input("Auto Frame Mesh Size (m)", min_value=0.05, max_value=5.0, value=0.25, step=0.05)

        if uploaded_dxf and st.button("Extract Data from DXF"):
            for key in ['dxf_parsed', 'adv_fea_data', 'num_loads_override']: 
                st.session_state.pop(key, None)
                
            st.session_state.adv_solved = False
            dxf_data = parse_dxf_to_data(uploaded_dxf.getvalue())
            
            if dxf_data:
                st.session_state.dxf_parsed = dxf_data
                st.session_state.num_loads_override = 0 
                st.success("✅ DXF Parsed Successfully!")
            else: 
                st.error("❌ Failed to extract meaningful data.")

        dxf_data = st.session_state.get('dxf_parsed', None)
        num_loads_def = st.session_state.get('num_loads_override', 0 if dxf_data else 1)

        c_in, c_plot = st.columns([1.2, 1.8])
        with c_in:
            st.markdown("### 1. Supports & Base System")
            def_supp_count = len(dxf_data['supports']) if dxf_data else 2
            num_base_sups = st.number_input("Count of Point Supports", 0, 50, def_supp_count, on_change=reset_adv_state)
            
            base_sups = []
            for i in range(int(num_base_sups)):
                sp1, sp2 = st.columns(2)
                def_sx = float(i*2.0)
                def_sy = 0.0
                dxf_seg_idx = None
                dxf_s_dist = None
                
                if dxf_data and i < len(dxf_data['supports']):
                    def_sx = dxf_data['supports'][i].get('x', 0.0)
                    def_sy = dxf_data['supports'][i].get('y', 0.0)
                    dxf_seg_idx = dxf_data['supports'][i].get('seg_idx')
                    dxf_s_dist = dxf_data['supports'][i].get('s_dist')
                    
                sx = sp1.number_input(f"Sup J{i+1} X (m)", value=float(def_sx), format="%.5f", on_change=reset_adv_state, key=f"sx_{i}")
                styp = sp2.selectbox(f"Sup J{i+1} Type", ["Hinged", "Roller", "Fixed"], key=f"sp_{i}", on_change=reset_adv_state)
                sup_dict = {'x': sx, 'y': def_sy, 'type': styp}
                
                if dxf_seg_idx is not None and abs(sx - def_sx) < 1e-4:
                    sup_dict['seg_idx'] = dxf_seg_idx
                    sup_dict['s_dist'] = dxf_s_dist
                base_sups.append(sup_dict)

            st.markdown("### 2. Segments")
            num_segs = st.number_input("Number of Segments", min_value=1, max_value=50, value=len(dxf_data['segments']) if dxf_data else 1, on_change=reset_adv_state)
            seg_choices = [f"S{i+1}" for i in range(int(num_segs))]
            segments = []
            
            for i in range(int(num_segs)):
                with st.expander(f"⚙️ Segment S{i+1}", expanded=(num_segs<3)):
                    s_type_idx = 0
                    def_L = 3.0
                    def_ang = 60.0
                    def_R = 5.0
                    def_S = 3.0
                    def_smooth = True
                    dir_crv_idx = 0
                    is_dxf = False
                    
                    abs_p1 = None
                    abs_p2 = None
                    abs_c = None
                    abs_r = None
                    abs_sa = None
                    abs_ea = None
                    sweep = None
                    
                    if dxf_data and i < len(dxf_data['segments']):
                        d_seg = dxf_data['segments'][i]
                        is_dxf = d_seg.get('is_dxf', False)
                        s_type_raw = d_seg.get('Shape Type', d_seg.get('type', 'Straight Line'))
                        
                        if s_type_raw == 'Straight Line':
                            s_type_idx = 0
                            def_L = d_seg.get('L', 3.0)
                            def_ang = d_seg.get('start_angle', 60.0)
                            def_smooth = d_seg.get('smooth', False)
                            abs_p1 = d_seg.get('abs_p1')
                            abs_p2 = d_seg.get('abs_p2')
                            
                        elif s_type_raw == 'Curve (Arc & Radius)':
                            s_type_idx = 1
                            def_R = d_seg.get('Radius (R) (m)', 5.0)
                            def_S = d_seg.get('L', 3.0)
                            def_ang = d_seg.get('start_angle', 0.0)
                            def_smooth = d_seg.get('smooth', False)
                            abs_c = d_seg.get('abs_c')
                            abs_r = d_seg.get('abs_r')
                            abs_sa = d_seg.get('abs_sa')
                            abs_ea = d_seg.get('abs_ea')
                            sweep = d_seg.get('sweep')
                            
                    if is_dxf:
                        st.success(f"🔒 DXF Locked: {s_type_raw}")
                        s_type = s_type_raw
                        L = def_L
                        kappa = d_seg.get('kappa', 0.0)
                        smooth = False
                        start_angle = def_ang
                    else:
                        s_type = st.radio(f"Shape Type (S{i+1})", ["Straight Line", "Curve (Arc)"], index=s_type_idx, key=f"t_{i}", horizontal=True, on_change=reset_adv_state)
                        smooth = True
                        start_angle = 0.0
                        
                        if i == 0: 
                            start_angle = st.number_input("Starting Angle (°)", value=float(def_ang), key=f"sa_{i}", on_change=reset_adv_state)
                            smooth = False
                        else:
                            smooth = st.checkbox(f"Smooth S{i+1}", value=def_smooth, key=f"sm_{i}", on_change=reset_adv_state)
                            if not smooth: 
                                start_angle = st.number_input(f"New Angle S{i+1} (°)", value=float(def_ang), key=f"sa_{i}", on_change=reset_adv_state)
                                
                        if s_type == "Straight Line": 
                            L = st.number_input(f"L (m)", value=float(def_L), key=f"l_{i}", on_change=reset_adv_state)
                            kappa = 0.0
                        else:
                            r_val = st.number_input(f"R (m)", value=float(def_R), key=f"r_{i}", on_change=reset_adv_state)
                            L = st.number_input(f"Arc L (m)", value=float(def_S), key=f"al_{i}", on_change=reset_adv_state)
                            kappa = 1.0/r_val
                    
                    seg_info = {
                        'type': s_type, 
                        'Shape Type': s_type, 
                        'L': L, 
                        'kappa': kappa, 
                        'smooth': smooth, 
                        'start_angle': start_angle
                    }
                    
                    if is_dxf:
                        seg_info['is_dxf'] = True
                        if s_type == "Straight Line": 
                            seg_info['abs_p1'] = abs_p1
                            seg_info['abs_p2'] = abs_p2
                        else: 
                            seg_info['abs_c'] = abs_c
                            seg_info['abs_r'] = abs_r
                            seg_info['abs_sa'] = abs_sa
                            seg_info['abs_ea'] = abs_ea
                            seg_info['sweep'] = sweep
                            
                    segments.append(seg_info)

            st.markdown("### 3. Properties")
            sec_list = list(SECTIONS_DB.keys()) if SECTIONS_DB else ["Soldier U100"]
            master_sec_name = st.selectbox("Master Section", sec_list, on_change=reset_adv_state)
            master_raw = SECTIONS_DB.get(master_sec_name, {})
            master_props = {
                'name': master_sec_name, 
                'E': master_raw.get('E', 2100.0), 
                'A': master_raw.get('A', 0.0034), 
                'I': master_raw.get('I', 412.0), 
                'Mall': master_raw.get('Mall', 13.1), 
                'Qall': master_raw.get('Qall', 100.8)
            }
            seg_sections = [master_props for _ in range(int(num_segs))]

            st.markdown("### 4. Applied Loads")
            num_loads = st.number_input("Count of Loads", 0, 30, num_loads_def, on_change=reset_adv_state)
            combined_loads = []
            
            for i in range(int(num_loads)):
                with st.expander(f"📥 Load {i+1}", expanded=(i==0)):
                    l_type = st.selectbox("Type", ["Uniform", "Point Load"], key=f"ld_t_{i}")
                    l_dir = st.selectbox("Direction", ["Global Z", "Global X", "Local Z"], key=f"ld_d_{i}")
                    s_choice = st.selectbox("Segment", seg_choices, key=f"ld_single_{i}")
                    s_idx_num = int(s_choice[1:]) - 1
                    w1 = st.number_input("Value W1", value=-15.0, key=f"ld_w1_{i}")
                    
                    combined_loads.append({
                        'seg_idx': s_idx_num, 
                        'type': l_type, 
                        'dir': l_dir, 
                        'start': 0.0, 
                        'end': float(segments[s_idx_num].get('L', 0.0)), 
                        'w1': w1, 
                        'w2': w1
                    })

            st.markdown("### 5. Struts")
            num_struts = st.number_input("Count of Struts", 0, 50, len(dxf_data['struts']) if dxf_data else 1, on_change=reset_adv_state)
            struts_data = []
            for i in range(int(num_struts)):
                struts_data.append({
                    'seg_idx': 0, 
                    's_dist': 1.0, 
                    'gx': 3.0, 
                    'gy': 0.0, 
                    'sec': "PPH 353"
                })

        nodes, elements, nodal_loads, display_nodes, supports_list, seg_starts = build_chain_mesh(segments, seg_sections, combined_loads, struts_data, None, base_sups, {'type': 'Hinged', 'angle': 0.0}, mesh_size=auto_mesh_size)

        with c_plot:
            st.markdown("### Live Geometry")
            live_img = draw_live_preview(nodes, elements, supports_list, seg_sections, combined_loads, segments, seg_starts)
            st.image(live_img, use_container_width=True)

        if st.button("🚀 Run Classic Analysis", type="primary"):
            U, R = solve_fea_engine(nodes, elements, nodal_loads, supports_list)
            st.session_state.adv_fea_data = {
                'U': U, 
                'R': R, 
                'nodes': nodes, 
                'elements': elements, 
                'display_nodes': display_nodes,
                'supports_list': supports_list, 
                'seg_sections': seg_sections,
                'loads_data': combined_loads, 
                'segments': segments, 
                'seg_starts': seg_starts
            }
            st.session_state.adv_solved = True

    # =======================================================
    # COMMON RENDER LOGIC (FOR BOTH INTERACTIVE & CLASSIC)
    # =======================================================
    if getattr(st.session_state, 'adv_solved', False):
        st.markdown("---")
        st.markdown("### 🎛️ Analysis Results & Diagrams")
        fea_data = st.session_state.adv_fea_data
        
        with st.expander("⚙️ Diagram Scale Controls", expanded=True):
            c_s1, c_s2, c_s3 = st.columns(3)
            sc_n = c_s1.slider("Axial Scale", 0.001, 0.100, 0.015, step=0.001)
            sc_v = c_s2.slider("Shear Scale", 0.001, 0.100, 0.015, step=0.001)
            sc_m = c_s3.slider("Moment Scale", 0.001, 0.100, 0.015, step=0.001)
            
        img_bufs = plot_sap2000_diagrams(
            fea_data['nodes'], 
            fea_data['elements'], 
            fea_data['R'], 
            {'N': sc_n, 'V': sc_v, 'M': sc_m}, 
            fea_data['display_nodes'], 
            fea_data['supports_list'], 
            fea_data['seg_sections'], 
            loads=fea_data['loads_data'], 
            segments=fea_data.get('segments'), 
            seg_starts=fea_data.get('seg_starts')
        )

        c_p1, c_p2, c_p3 = st.columns(3)
        c_p1.image(img_bufs['React'], use_container_width=True)
        c_p2.image(img_bufs['V'], use_container_width=True)
        c_p3.image(img_bufs['M'], use_container_width=True)
        
        st.markdown("### 📊 Safety Summary")
        safety_data = []
        for i, sec in enumerate(fea_data['seg_sections']):
            max_m = 0.0
            max_v = 0.0
            for el in fea_data['elements']:
                if el.get('group') == 'segment':
                    max_m = max(max_m, np.max(np.abs(el.get('internal', {}).get('M', [0]))))
                    max_v = max(max_v, np.max(np.abs(el.get('internal', {}).get('V', [0]))))
            
            s_status = "SAFE" if max_m <= sec['Mall'] and max_v <= sec['Qall'] else "UNSAFE"
            safety_data.append({
                "Component": f"S{i+1}", 
                "Force Type": "Bending & Shear", 
                "Actual": f"M={max_m:.1f}, V={max_v:.1f}", 
                "Allowable": f"M={sec['Mall']:.1f}, V={sec['Qall']:.1f}", 
                "Status": s_status
            })
            
        st.table(pd.DataFrame(safety_data))
        fea_data['safety_df'] = safety_data
        fea_data['img_bufs'] = img_bufs
        
        doc_out = generate_chain_report(fea_data)
        st.download_button(
            "⬇️ Download Calculation Sheet (Word)", 
            data=doc_out.getvalue(), 
            file_name="Advanced_Shape_Calculation_Sheet.docx"
        )

if __name__ == "__main__":
    render_advanced_shape_module()
